import re
import os
from docx import Document
import NON_OCR
import subprocess
import shutil
import uuid

def main(status, important_words_list, file_path, output_path):
    '''
    status : 파일 형식이 문자열로 담겨 있는 변수, ex '.txt'
    important_words_list : GPT에게서 받은 중요 키워드 리스트
    file_path : 원본 파일 경로
    output_path : 최종 결과물 저장 경로
    '''
    if status == '.txt':
        extracted_text = NON_OCR.read_file_with_detected_encoding(file_path)
        txt(extracted_text, important_words_list, output_path)
    
    elif status == '.docx':
        docx(file_path, output_path, important_words_list)
        
    elif status == '.doc':
        doc(file_path, output_path, important_words_list)
        
    


# txt파일에 대해 빈칸 생성

def txt(text, keywords, output_path):
    '''
    text : extracted_text
    keywords : important_words_list
    output_path : 저장할 파일 경로
    '''
    # 모든 키워드의 스팬 수집
    spans = []
    for keyword in keywords:
        for match in re.finditer(keyword, text):
            spans.append((match.start(), match.end()))

    # 스팬 정렬 후 병합
    spans.sort()  # 시작 위치 기준으로 정렬
    merged_spans = []
    for start, end in spans:
        if not merged_spans or start >= merged_spans[-1][1]:  # 겹치지 않으면 추가
            merged_spans.append((start, end))
        else:  # 겹치면 기존 스팬 확장
            merged_spans[-1] = (merged_spans[-1][0], max(merged_spans[-1][1], end))

    # 병합된 스팬을 기준으로 치환
    result = []
    last_index = 0
    for start, end in merged_spans:
        result.append(text[last_index:start])  # 이전 텍스트 추가
        blank_length = end - start
        result.append("(" + "_" * blank_length + ")")  # 길이에 비례한 공백
        last_index = end
    result.append(text[last_index:])  # 마지막 남은 텍스트 추가

    modified_text = "".join(result)
    
    # 저장하기
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(modified_text)

    print('최종 결과물 저장 완료')



# 문자열 치환 함수
def docx(input_path, output_path, important_words_list):
    """
    .docx 파일에서 문자열을 치환하고 새로운 파일로 저장합니다.

    Args:
        input_path (str): 입력 파일 경로
        output_path (str): 출력 파일 경로
        important_words_list (list): 중요 단어 리스트
    """
    # 리스트가 비어 있으면 작업 중단
    if not important_words_list:
        print("중요 단어 리스트가 비어 있습니다. 작업을 중단합니다.")
        return

    # {old_str: new_str} 으로 구성된 replacements 딕셔너리 생성
    replacements = {
        keyword: "(__" + "_" * max(len(keyword) - 2, 0) + "__)" for keyword in important_words_list
    }
    
    # 입력 파일 확인
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"입력 파일을 찾을 수 없습니다: {input_path}")

    # 문서 열기
    try:
        doc = Document(input_path)
    except Exception as e:
        raise Exception(f"문서를 열 수 없습니다: {e}")

    # 단락 내 문자열 치환
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)

    # 표 내 문자열 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)

    # 수정된 문서 저장
    try:
        doc.save(output_path)
        print(f"치환된 문서를 저장했습니다: {output_path}")
    except Exception as e:
        raise Exception(f"문서를 저장할 수 없습니다: {e}")



def doc(file_path, output_path, important_words_list):
    """
    .doc 파일을 .docx로 변환한 후, 문자열을 치환하고 저장합니다.

    Args:
        file_path (str): 입력 .doc 파일 경로
        output_path (str): 출력 .docx 파일 경로
        important_words_list (list): 중요 단어 리스트
    """
    try:
        # LibreOffice 실행 파일의 경로 (필요 시 수정)
        libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

        # 입력 파일 확인
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"입력 파일을 찾을 수 없습니다: {file_path}")

        # 고유한 임시 출력 디렉토리 생성
        temp_dir = f"temp_output_{uuid.uuid4().hex}"
        os.makedirs(temp_dir, exist_ok=True)

        # LibreOffice CLI를 사용하여 .doc -> .docx 변환
        subprocess.run([
            libreoffice_path, 
            "--headless", 
            "--convert-to", "docx", 
            "--outdir", temp_dir, 
            file_path
        ], check=True)

        # 변환된 .docx 파일 경로
        base_name = os.path.basename(file_path)
        converted_path = os.path.join(temp_dir, os.path.splitext(base_name)[0] + ".docx")

        # 변환된 파일이 실제로 존재하는지 확인
        if not os.path.exists(converted_path):
            raise FileNotFoundError(f"LibreOffice 변환 결과 파일을 찾을 수 없습니다: {converted_path}")

        # docx 치환 작업 수행
        docx(converted_path, output_path, important_words_list)

    except subprocess.CalledProcessError as e:
        print(f"LibreOffice 변환 중 오류가 발생했습니다: {e}")
    except Exception as e:
        print(f"예기치 않은 오류가 발생했습니다: {e}")
    finally:
        # 임시 디렉토리 삭제
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

