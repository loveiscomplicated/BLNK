import os
from docx import Document
from charset_normalizer import detect
import olefile
import zipfile
from lxml import etree
import subprocess
import os


    



def read_file_with_detected_encoding(file_path):
    """
    파일의 인코딩을 감지하여 내용을 읽어옵니다.

    Args:
        file_path (str): 읽을 파일의 경로.

    Returns:
        str: 파일 내용.
    """
    try:
        with open(file_path, "rb") as file:
            raw_data = file.read()
            detected = detect(raw_data)
            encoding = detected["encoding"]

        with open(file_path, "r", encoding=encoding) as file:
            content = file.read()
    except Exception as e:
        return f"Error reading .txt file: {e}"
    return content



def docx2text(file_path):
    try:
        # .docx 파일 열기
        doc = Document(file_path)
        
        # 모든 문단 텍스트를 결합하여 반환
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return content
    except Exception as e:
        return f"Error reading .docx file: {e}"
    




def read_doc(file_path):
    """
    LibreOffice CLI를 사용하여 .doc 파일을 .docx 파일로 변환하고 내용을 반환합니다.

    Args:
        file_path (str): 변환할 .doc 파일의 경로.

    Returns:
        str: 변환된 파일의 내용 또는 에러 메시지.
    """
    try:
        # LibreOffice 실행 파일의 경로를 명시적으로 지정 @@@@@@@@ 나중에 꼭 바꾸기 @@@@@@@@
        libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

        # 임시 출력 디렉토리 생성
        output_dir = "temp_output"
        os.makedirs(output_dir, exist_ok=True)

        # LibreOffice CLI를 실행하여 파일 변환
        subprocess.run([
            libreoffice_path, 
            "--headless", 
            "--convert-to", "docx", 
            "--outdir", output_dir, 
            file_path
        ], check=True)

        # 변환된 파일 경로 생성
        base_name = os.path.basename(file_path)
        output_file = os.path.join(output_dir, os.path.splitext(base_name)[0] + ".docx")

        if os.path.exists(output_file):
            # 변환된 파일 읽기
            from docx import Document
            doc = Document(output_file)
            content = "\n".join([para.text for para in doc.paragraphs])

            # 임시 파일 삭제
            os.remove(output_file)
            os.rmdir(output_dir)

            return content
        else:
            return "Error: Output file not found."

    except subprocess.CalledProcessError as e:
        return f"Error during conversion: {e}"
    except Exception as e:
        return f"Unexpected error: {e}"



def hwp2text(file_path):
    try:
        # .hwp 파일 열기
        f = olefile.OleFileIO(file_path)
        
        # 파일 내용 가져오기
        encoded_text = f.openstream('PrvText').read() 
        decoded_text = encoded_text.decode('utf-16')  
        
        return decoded_text
    except Exception as e:
        return f"Error reading .hwp file: {e}"


def list_hwpx_files(file_path):
    """
    .hwpx 파일 내부의 모든 파일 목록을 확인합니다.
    """
    with zipfile.ZipFile(file_path, 'r') as z:
        return z.namelist()

def extract_hwpx_text(file_path):
    """
    동적 네임스페이스 처리를 통해 .hwpx 파일에서 텍스트를 추출합니다.
    """
    with zipfile.ZipFile(file_path, 'r') as z:
        # 텍스트가 포함된 것으로 예상되는 콘텐츠 파일 식별
        content_files = [f for f in z.namelist() if 'Contents' in f]
        text_data = []

        for file in content_files:
            with z.open(file) as content:
                tree = etree.parse(content)
                root = tree.getroot()

                # 동적으로 네임스페이스 추출
                namespaces = {k: v for k, v in root.nsmap.items() if k}
                #print("발견된 네임스페이스:", namespaces)

                # 네임스페이스와 함께 XPath를 사용하여 텍스트 추출
                for paragraph in tree.xpath('//hp:t', namespaces=namespaces):
                    if paragraph.text:
                        text_data.append(paragraph.text)

        return "\n".join(text_data)
    
    



if __name__ == '__main__':
    #file_path = "C:/Users/LG/miniconda3/envs/BLNK/codes/sample_doc.doc"
    #content = read_doc(file_path)
    #print(content)
    
    #file_path = "C:/Users/LG/miniconda3/envs/BLNK/codes/sample_docx.docx"
    #content = docx2text(file_path)
    #print(content)

    #file_path = "C:/Users/LG/miniconda3/envs/BLNK/codes/sample_txt.txt"
    #content = read_file_with_detected_encoding(file_path)
    #print(content) 
    file_path = 'C:/Users/LG/miniconda3/envs/BLNK/codes/check/과연.hwpx'  # 사용할 .hwpx 파일 경로로 교체하세요

    # 단계 1: .hwpx 압축 파일 내부의 모든 파일 목록 출력
    #print(".hwpx 파일의 내부 파일:")
    #print(list_hwpx_files(file_path))

    # 단계 2: .hwpx 파일에서 텍스트 추출
    print("\n추출된 텍스트:")
    extracted_text = extract_hwpx_text(file_path)
    print(extracted_text)